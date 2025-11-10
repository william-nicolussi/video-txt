# NotifyTime.pyw
# Whisper Study Suite with bilingual UI (EN/IT), help popups, optional diarization,
# study-friendly exports (Markdown, SRT, VTT, DOCX, CSV Q&A) and LLM JSON.
# Default: saves ONLY raw .txt unless extra exports are selected.

import os, sys, math, time, tempfile, threading, queue, subprocess, json, csv, re, textwrap, pathlib
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import timedelta
from collections import Counter, defaultdict

from plyer import notification
import whisper

# ---------- Optional diarization ----------
_HAS_PYANNOTE = False
try:
    from pyannote.audio import Pipeline
    from pyannote.core import Annotation, Segment
    _HAS_PYANNOTE = True
except Exception:
    pass

# ---------- Optional DOCX export ----------
_HAS_DOCX = False
try:
    from docx import Document
    from docx.shared import Pt
    _HAS_DOCX = True
except Exception:
    pass

# ================= i18n =================
LANGS = ["en", "it"]
I18N = {
    "en": {
        "APP_NAME": "Whisper Study Suite",
        "Language UI": "Language",
        # Groups / labels
        "Media files": "Media files",
        "Add files…": "Add files…",
        "Clear": "Clear",
        "No files selected": "No files selected",
        "Output folder": "Output folder",
        "Browse…": "Browse…",
        "Transcription": "Transcription",
        "Model": "Model",
        "Chunk length (s)": "Chunk length (s)",
        "Save every N chunks": "Save every N chunks",
        "Language (blank=auto)": "Language (blank=auto)",
        "Known speakers (0=auto, 1=single)": "Known speakers (0=auto, 1=single)",
        "Silence → new paragraph (s)": "Silence → new paragraph (s)",
        "Pre-clean audio (denoise)": "Pre-clean audio (denoise)",
        "Segments": "Segments",
        "Max segment length (s, 0=off)": "Max segment length (s, 0=off)",
        "Concept Highlighter & Topic Markers": "Concept Highlighter & Topic Markers",
        "Enable keyword highlighting + glossary": "Enable keyword highlighting + glossary",
        "Top-N keywords": "Top-N keywords",
        "Topic markers (comma-separated)": "Topic markers (comma-separated)",
        "Optional exports (raw .txt is always saved)": "Optional exports (raw .txt is always saved)",
        "TXT refined": "TXT refined",
        "SRT": "SRT",
        "VTT": "VTT",
        "Markdown": "Markdown",
        "DOCX": "DOCX",
        "CSV Q&A": "CSV Q&A",
        "JSON for LLM": "JSON for LLM",
        "Start": "Start",
        "Idle": "Idle",
        # Dialog titles/messages
        "What is this?": "What is this?",
        "A job is already running.": "A job is already running.",
        "Select at least one media file.": "Select at least one media file.",
        "ffmpeg not found in PATH.": "ffmpeg not found in PATH.",
        # Help texts
        "help_model": "ASR model size. Larger models are slower but more accurate. Suggestion: 'small' for English lectures; 'medium' for harder audio.",
        "help_chunk": "Split audio into fixed-length chunks before transcription. Longer chunks reduce overhead; shorter chunks checkpoint progress more often.",
        "help_save_every": "Write the running raw .txt to disk every N chunks so you never lose much progress.",
        "help_language": "Force a language (e.g., 'en'). Leave blank for automatic detection.",
        "help_known_speakers": "0: try automatic speaker diarization (needs pyannote + HF token). 1: single speaker (skip diarization). >1: multi-speaker if available.",
        "help_silence_gap": "Start a new paragraph when silence between segments exceeds this value (seconds).",
        "help_preclean": "Apply a light denoise/EQ via ffmpeg to improve punctuation and segmentation on noisy audio.",
        "help_max_segment": "Re-slice paragraphs so each block is at most this long (seconds). Useful for study and precise citations.",
        "help_highlight": "Detect frequent terms, bold them in refined outputs, and build a small glossary with timestamps.",
        "help_markers": "Comma-separated words that signal topic changes (e.g., slide, theorem, lemma). Used to tag each block’s 'topic' in JSON.",
        "help_txt_refined": "Readable TXT with speaker, timestamps, and paragraphs.",
        "help_srt": "SubRip subtitles with timestamps and speaker labels.",
        "help_vtt": "WebVTT subtitles for web players.",
        "help_md": "Markdown with headings and optional glossary.",
        "help_docx": "Formatted Word document with optional glossary.",
        "help_csvqa": "Simple Q&A CSV extracted from informative sentences.",
        "help_llm_json": "Structured JSON for LLM prompts: topic, start/end, speaker, text, keywords."
    },
    "it": {
        "APP_NAME": "Whisper Study Suite",
        "Language UI": "Lingua",
        "Media files": "File multimediali",
        "Add files…": "Aggiungi file…",
        "Clear": "Pulisci",
        "No files selected": "Nessun file selezionato",
        "Output folder": "Cartella di output",
        "Browse…": "Sfoglia…",
        "Transcription": "Trascrizione",
        "Model": "Modello",
        "Chunk length (s)": "Lunghezza chunk (s)",
        "Save every N chunks": "Salva ogni N chunk",
        "Language (blank=auto)": "Lingua (vuoto=auto)",
        "Known speakers (0=auto, 1=single)": "Parlanti noti (0=auto, 1=singolo)",
        "Silence → new paragraph (s)": "Silenzio → nuovo paragrafo (s)",
        "Pre-clean audio (denoise)": "Pre-pulizia audio (denoise)",
        "Segments": "Segmenti",
        "Max segment length (s, 0=off)": "Lunghezza max segmento (s, 0=off)",
        "Concept Highlighter & Topic Markers": "Evidenziatore concetti & Marker di argomento",
        "Enable keyword highlighting + glossary": "Abilita evidenziazione parole chiave + glossario",
        "Top-N keywords": "Top-N parole chiave",
        "Topic markers (comma-separated)": "Marker di argomento (separati da virgola)",
        "Optional exports (raw .txt is always saved)": "Export opzionali (il .txt grezzo è sempre salvato)",
        "TXT refined": "TXT pulito",
        "SRT": "SRT",
        "VTT": "VTT",
        "Markdown": "Markdown",
        "DOCX": "DOCX",
        "CSV Q&A": "CSV Q&A",
        "JSON for LLM": "JSON per LLM",
        "Start": "Avvia",
        "Idle": "In attesa",
        "What is this?": "Che cos’è?",
        "A job is already running.": "Un processo è già in esecuzione.",
        "Select at least one media file.": "Seleziona almeno un file.",
        "ffmpeg not found in PATH.": "ffmpeg non è nel PATH.",
        "help_model": "Dimensione del modello ASR. Più grande = più lento ma più accurato.",
        "help_chunk": "Divide l’audio in chunk di lunghezza fissa prima della trascrizione.",
        "help_save_every": "Scrive il .txt grezzo ogni N chunk, per non perdere avanzamento.",
        "help_language": "Forza la lingua (es. 'en'); vuoto = auto.",
        "help_known_speakers": "0: diarization automatica (richiede pyannote + token HF). 1: singolo parlante. >1: multi-parlante se disponibile.",
        "help_silence_gap": "Nuovo paragrafo quando il silenzio supera questo valore (secondi).",
        "help_preclean": "Denoise/EQ leggero via ffmpeg per migliorare punteggiatura e segmentazione.",
        "help_max_segment": "Ritaglia i paragrafi per avere blocchi di durata massima (secondi).",
        "help_highlight": "Rileva termini frequenti, li mette in grassetto e crea un glossario con timestamp.",
        "help_markers": "Parole che segnalano cambio argomento (es. slide, teorema, lemma). Usate per tag nel JSON.",
        "help_txt_refined": "TXT leggibile con parlante, timestamp e paragrafi.",
        "help_srt": "Sottotitoli SubRip con timestamp e parlante.",
        "help_vtt": "Sottotitoli WebVTT per web player.",
        "help_md": "Markdown con intestazioni e glossario opzionale.",
        "help_docx": "Documento Word formattato con glossario opzionale.",
        "help_csvqa": "CSV Q&A da frasi informative.",
        "help_llm_json": "JSON strutturato per LLM: topic, tempi, parlante, testo, keyword."
    }
}

_cfg_path = pathlib.Path.home() / ".whisper_study_suite.json"
def load_ui_lang():
    try:
        if _cfg_path.exists():
            return json.loads(_cfg_path.read_text()).get("ui_lang", "en")
    except Exception:
        pass
    return "en"

def save_ui_lang(lang):
    try:
        _cfg_path.write_text(json.dumps({"ui_lang": lang}))
    except Exception:
        pass

_current_lang = load_ui_lang()

def tr(key: str) -> str:
    return I18N.get(_current_lang, I18N["en"]).get(key, key)

# ================= core helpers =================
EN_STOPWORDS = set("""
a about above after again against all am an and any are aren't as at be because been before being
below between both but by can't cannot could couldn't did didn't do does doesn't doing don't down
during each few for from further had hadn't has hasn't have haven't having he he'd he'll he's her here
here's hers herself him himself his how how's i i'd i'll i'm i've if in into is isn't it it's its
itself let's me more most mustn't my myself no nor not of off on once only or other ought our ours
ourselves out over own same shan't she she'd she'll she's should shouldn't so some such than that
that's the their theirs them themselves then there there's these they they'd they'll they're they've
this those through to too under until up very was wasn't we we'd we'll we're we've were weren't what
what's when when's where where's which while who who's whom why why's with won't would wouldn't you
you'd you'll you're you've your yours yourself yourselves
""".split())

def check_ffmpeg():
    try:
        subprocess.run(["ffmpeg", "-version"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, check=True)
        return True
    except Exception:
        return False

def extract_audio_from_media(media_path, sample_rate=16000, mono=True, denoise=False, log=None):
    base = os.path.splitext(os.path.basename(media_path))[0]
    tmpdir = tempfile.mkdtemp()
    wav_path = os.path.join(tmpdir, base + ".wav")
    if log: log(f"Extracting audio: {media_path}")
    cmd = ["ffmpeg", "-i", media_path, "-vn", "-acodec", "pcm_s16le", "-ar", str(sample_rate)]
    if mono: cmd += ["-ac", "1"]
    if denoise:
        cmd += ["-af", "highpass=f=80,lowpass=f=8000,arnndn=m=rnnoise"]
    cmd += [wav_path, "-y"]
    subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return wav_path

def split_audio(audio_path, chunk_length, log=None):
    if log: log(f"Splitting audio into {chunk_length}s chunks…")
    d = tempfile.mkdtemp()
    pat = os.path.join(d, "chunk_%03d.wav")
    cmd = ["ffmpeg", "-i", audio_path, "-f", "segment", "-segment_time", str(chunk_length), "-c", "copy", pat, "-y"]
    subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return sorted([os.path.join(d, f) for f in os.listdir(d) if f.endswith(".wav")])

def _fmt_ts_srt(t: float) -> str:
    td = timedelta(seconds=max(0, t))
    h = int(td.total_seconds() // 3600)
    m = int((td.total_seconds() % 3600) // 60)
    s = int(td.total_seconds() % 60)
    ms = int((td.total_seconds() - int(td.total_seconds())) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

def _fmt_ts_vtt(t: float) -> str:
    td = timedelta(seconds=max(0, t))
    h = int(td.total_seconds() // 3600)
    m = int((td.total_seconds() % 3600) // 60)
    s = int(td.total_seconds() % 60)
    ms = int((td.total_seconds() - int(td.total_seconds())) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"

def normalize_sentence(text: str) -> str:
    t = text.strip()
    if not t: return t
    if t[0].islower(): t = t[0].upper() + t[1:]
    if not t.endswith(('.', '!', '?')): t += '.'
    t = t.replace(" ,", ",").replace(" .", ".").replace(" !", "!").replace(" ?", "?")
    t = re.sub(r"\s+", " ", t)
    return t.strip()

# ---------- diarization glue ----------
def _majority_speaker(st: float, et: float, diar: "Annotation") -> str:
    try:
        seg = Segment(st, et)
        overlaps = diar.crop(seg, mode="intersection")
        speaker_dur = {}
        for s, _, label in overlaps.itertracks(yield_label=True):
            dur = max(0.0, min(et, s.end) - max(st, s.start))
            speaker_dur[label] = speaker_dur.get(label, 0.0) + dur
        if not speaker_dur:
            return "Speaker ?"
        return max(speaker_dur.items(), key=lambda kv: kv[1])[0]
    except Exception:
        return "Speaker ?"

def align_whisper_with_diarization(wh_segments, diar):
    for seg in wh_segments:
        seg["speaker"] = _majority_speaker(seg["start"], seg["end"], diar)
    return wh_segments

# ---------- grouping & segmentation ----------
def paragraphize(segments, silence_gap):
    if not segments: return []
    groups = []
    cur = {"speaker": segments[0].get("speaker","Speaker"), "start": segments[0]["start"], "end": segments[0]["end"], "texts":[segments[0]["text"].strip()]}
    prev_end = segments[0]["end"]
    for s in segments[1:]:
        gap = s["start"] - prev_end
        same_speaker = s.get("speaker","Speaker") == cur["speaker"]
        if (not same_speaker) or (gap >= silence_gap):
            cur["text"] = " ".join(cur["texts"]).strip()
            groups.append(cur)
            cur = {"speaker": s.get("speaker","Speaker"), "start": s["start"], "end": s["end"], "texts":[s["text"].strip()]}
        else:
            cur["texts"].append(s["text"].strip())
            cur["end"] = s["end"]
        prev_end = s["end"]
    cur["text"] = " ".join(cur["texts"]).strip()
    groups.append(cur)
    return groups

def reslice_groups_by_window(groups, max_seconds):
    if max_seconds <= 0: return groups
    out = []
    for g in groups:
        start, end, text, speaker = g["start"], g["end"], g["text"], g["speaker"]
        if end - start <= max_seconds:
            out.append(g); continue
        words = text.split()
        if not words:
            out.append(g); continue
        total_dur = end - start
        avg_per_word = total_dur / max(1, len(words))
        window_words = max(1, int(max_seconds / avg_per_word))
        i = 0; cur_start = start
        while i < len(words):
            chunk_words = words[i:i+window_words]
            chunk_text = " ".join(chunk_words)
            chunk_dur = len(chunk_words) * avg_per_word
            out.append({
                "speaker": speaker,
                "start": cur_start,
                "end": min(end, cur_start + chunk_dur),
                "text": normalize_sentence(chunk_text)
            })
            cur_start += chunk_dur
            i += window_words
    return out

# ---------- concepts / glossary / topics ----------
def tokenize_for_keywords(text: str):
    tokens = re.findall(r"[A-Za-z][A-Za-z\-']+", text.lower())
    tokens = [t for t in tokens if t not in EN_STOPWORDS and len(t) > 2]
    return tokens

def extract_keywords(groups, top_n=25):
    corpus = " ".join(g["text"] for g in groups)
    toks = tokenize_for_keywords(corpus)
    freq = Counter(toks)
    return [w for w, _ in freq.most_common(top_n)]

def annotate_keywords(groups, keywords):
    kw_set = set(k.lower() for k in keywords)
    annotated = []
    kw_hits = defaultdict(list)
    for g in groups:
        def repl(m):
            w = m.group(0)
            if w.lower() in kw_set:
                ts = _fmt_ts_srt(g['start'])[:-4]
                kw_hits[w.lower()].append(ts)
                return f"**{w}**"
            return w
        text = re.sub(r"[A-Za-z][A-Za-z\-']+", repl, g["text"])
        annotated.append({**g, "text": text})
    glossary = []
    for k in keywords:
        stamps = kw_hits.get(k.lower(), [])
        glossary.append({"term": k, "occurrences": len(stamps), "timestamps": stamps[:10]})
    return annotated, glossary

def infer_topics(groups, markers):
    topics = []
    current = "General"
    for g in groups:
        found = None
        for m in markers:
            if re.search(rf"\b{re.escape(m)}\b", g["text"], flags=re.IGNORECASE):
                found = m; break
        if found: current = found
        topics.append(current)
    return topics

# ---------- exports ----------
def save_raw_txt(fulltext, out_base, log=None):
    p = out_base + ".txt"
    with open(p, "w", encoding="utf-8") as f: f.write(fulltext)
    if log: log(f"Saved {os.path.basename(p)}")

def save_txt(groups, out_base, log=None):
    p = out_base + "_refined.txt"
    with open(p, "w", encoding="utf-8") as f:
        for g in groups:
            head = f"{g['speaker']} [{_fmt_ts_srt(g['start'])} – {_fmt_ts_srt(g['end'])}]"
            f.write(head + "\n" + g["text"] + "\n\n")
    if log: log(f"Saved {os.path.basename(p)}")

def save_srt(groups, out_base, log=None):
    p = out_base + "_refined.srt"
    with open(p, "w", encoding="utf-8") as f:
        for i, g in enumerate(groups, 1):
            f.write(f"{i}\n{_fmt_ts_srt(g['start'])} --> {_fmt_ts_srt(g['end'])}\n{g['speaker']}: {g['text']}\n\n")
    if log: log(f"Saved {os.path.basename(p)}")

def save_vtt(groups, out_base, log=None):
    p = out_base + "_refined.vtt"
    with open(p, "w", encoding="utf-8") as f:
        f.write("WEBVTT\n\n")
        for g in groups:
            f.write(f"{_fmt_ts_vtt(g['start'])} --> {_fmt_ts_vtt(g['end'])}\n{g['speaker']}: {g['text']}\n\n")
    if log: log(f"Saved {os.path.basename(p)}")

def save_markdown(groups, out_base, title, glossary=None, log=None):
    p = out_base + "_refined.md"
    with open(p, "w", encoding="utf-8") as f:
        f.write(f"# {title}\n\n")
        for g in groups:
            f.write(f"**{g['speaker']}** · {_fmt_ts_srt(g['start'])}–{_fmt_ts_srt(g['end'])}\n\n{g['text']}\n\n")
        if glossary:
            f.write("\n---\n\n## Glossary\n\n")
            for item in glossary:
                ts = ", ".join(item["timestamps"]) if item["timestamps"] else "-"
                f.write(f"- **{item['term']}** · {item['occurrences']} occ · {ts}\n")
    if log: log(f"Saved {os.path.basename(p)}")

def save_docx(groups, out_base, title, glossary=None, log=None):
    if not _HAS_DOCX:
        if log: log("python-docx not installed: skipping DOCX.")
        return
    p = out_base + "_refined.docx"
    doc = Document()
    doc.add_heading(title, level=1)
    for g in groups:
        head = f"{g['speaker']} · {_fmt_ts_srt(g['start'])}–{_fmt_ts_srt(g['end'])}"
        doc.add_paragraph(head).runs[0].bold = True
        para = doc.add_paragraph(g["text"])
        para.paragraph_format.space_after = Pt(6)
    if glossary:
        doc.add_page_break()
        doc.add_heading("Glossary", level=2)
        for item in glossary:
            line = f"{item['term']} — {item['occurrences']} occ"
            if item["timestamps"]:
                line += f" — {', '.join(item['timestamps'][:10])}"
            doc.add_paragraph(line)
    doc.save(p)
    if log: log(f"Saved {os.path.basename(p)}")

def save_csv_qa(groups, out_base, log=None):
    p = out_base + "_qa.csv"
    with open(p, "w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Question","Answer"])
        for g in groups:
            sentences = re.split(r'(?<=[.!?])\s+', g["text"])
            for s in sentences:
                s = s.strip()
                if 30 <= len(s) <= 180:
                    q = f"What does the lecturer explain at [{_fmt_ts_srt(g['start'])[:-4]}]?"
                    writer.writerow([q, s])
    if log: log(f"Saved {os.path.basename(p)}")

def save_llm_json(groups, out_base, topics, keywords, log=None):
    p = out_base + "_llm.json"
    payload = []
    for g, topic in zip(groups, topics):
        payload.append({
            "topic": topic,
            "start": g["start"],
            "end": g["end"],
            "speaker": g["speaker"],
            "text": g["text"],
            "keywords": keywords[:10]
        })
    with open(p, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    if log: log(f"Saved {os.path.basename(p)}")

# ================= transcription =================
def transcribe_one(model, media_path, chunk_length, save_every_chunks, language,
                   known_speakers, silence_gap, out_dir,
                   pre_clean_audio,
                   max_segment_seconds,
                   enable_keywords, topn_keywords,
                   markers_list,
                   export_txt_refined, export_srt, export_vtt, export_md, export_docx, export_csvqa, export_llm_json,
                   progress_cb, log):
    wav_path = extract_audio_from_media(media_path, denoise=pre_clean_audio, log=log)
    chunks = split_audio(wav_path, chunk_length, log=log)
    if not chunks:
        log("No chunks created. Check the media file.")
        return

    base_name = os.path.splitext(os.path.basename(media_path))[0]
    out_base = os.path.join(out_dir, base_name)
    start_time = time.time()

    diar = None
    if known_speakers == 0:
        if _HAS_PYANNOTE:
            token = os.environ.get("HUGGINGFACE_TOKEN", "").strip()
            if token:
                log("Running speaker diarization on full audio…")
                try:
                    pipe = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", use_auth_token=token)
                    diar = pipe(wav_path)
                except Exception as e:
                    log(f"Diarization failed: {e}. Continuing without diarization.")
            else:
                log("HUGGINGFACE_TOKEN not set. Skipping diarization.")
        else:
            log("pyannote not installed. Skipping diarization.")

    all_segments = []
    output_fulltext = ""
    total = len(chunks)
    for i, ch in enumerate(chunks):
        result = model.transcribe(
        ch,
        language=language if language else None,
        fp16=False,
        condition_on_previous_text=True   # o False, come preferisci
        )
        segs = result.get("segments", [])
        for s in segs:
            t = normalize_sentence(s.get("text",""))
            output_fulltext += t + "\n"
            all_segments.append({
                "start": float(s["start"]) + i*chunk_length,
                "end": float(s["end"]) + i*chunk_length,
                "text": t
            })
        if ((i + 1) % save_every_chunks == 0) or (i == total - 1):
            save_raw_txt(output_fulltext, out_base, log=log)
        elapsed = time.time() - start_time
        frac = (i + 1) / total
        eta = elapsed/frac - elapsed if frac > 0 else 0
        if progress_cb: progress_cb(frac, elapsed, eta)

    if known_speakers <= 1:
        for s in all_segments: s["speaker"] = "Speaker"
    else:
        if diar:
            all_segments = align_whisper_with_diarization(all_segments, diar)
        else:
            for s in all_segments: s["speaker"] = "Speaker"

    groups = paragraphize(all_segments, silence_gap)

    if max_segment_seconds > 0:
        groups = reslice_groups_by_window(groups, max_segment_seconds)

    glossary = None
    keywords = []
    if enable_keywords:
        keywords = extract_keywords(groups, top_n=topn_keywords)
        groups, glossary = annotate_keywords(groups, keywords)

    if any([export_txt_refined, export_srt, export_vtt, export_md, export_docx, export_csvqa, export_llm_json]):
        if export_txt_refined: save_txt(groups, out_base, log=log)
        if export_srt: save_srt(groups, out_base, log=log)
        if export_vtt: save_vtt(groups, out_base, log=log)
        if export_md: save_markdown(groups, out_base, title=base_name, glossary=glossary, log=log)
        if export_docx: save_docx(groups, out_base, title=base_name, glossary=glossary, log=log)
        if export_csvqa: save_csv_qa(groups, out_base, log=log)
        markers = [m.strip() for m in markers_list.split(",") if m.strip()] if markers_list else []
        topics = infer_topics(groups, markers) if markers else ["General"] * len(groups)
        if export_llm_json: save_llm_json(groups, out_base, topics, keywords, log=log)

    log(f"Completed: {out_base} (raw .txt always saved)")

# ================= GUI =================
def help_btn(parent, help_key):
    def show():
        messagebox.showinfo(tr("What is this?"), tr(help_key))
    b = ttk.Button(parent, text="?", width=2, command=show)
    return b

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.ui_lang = _current_lang
        self.title(tr("APP_NAME"))
        self.geometry("980x780"); self.minsize(940, 740)
        self.files = []; self.out_dir = os.path.expanduser("~")
        self.log_queue = queue.Queue(); self.worker_thread = None; self.model = None

        # Top language bar
        topbar = ttk.Frame(self)
        topbar.pack(fill="x", padx=10, pady=(10,0))
        ttk.Label(topbar, text=tr("Language UI")).pack(side="left")
        self.lang_var = tk.StringVar(value=self.ui_lang)
        ttk.Combobox(topbar, textvariable=self.lang_var, values=LANGS, width=6, state="readonly").pack(side="left", padx=6)
        ttk.Button(topbar, text="OK", command=self.on_change_lang).pack(side="left")

        # Media section
        self.lf_files = ttk.LabelFrame(self, text=tr("Media files")); self.lf_files.pack(fill="x", padx=10, pady=8)
        self.btn_add = ttk.Button(self.lf_files, text=tr("Add files…"), command=self.add_files); self.btn_add.pack(side="left", padx=6, pady=6)
        self.btn_clear = ttk.Button(self.lf_files, text=tr("Clear"), command=self.clear_files); self.btn_clear.pack(side="left", padx=6, pady=6)
        self.lbl_files = ttk.Label(self.lf_files, text=tr("No files selected")); self.lbl_files.pack(side="left", padx=10, pady=6)

        # Output
        self.lf_out = ttk.LabelFrame(self, text=tr("Output folder")); self.lf_out.pack(fill="x", padx=10, pady=8)
        self.out_var = tk.StringVar(value=self.out_dir)
        ttk.Entry(self.lf_out, textvariable=self.out_var).pack(side="left", fill="x", expand=True, padx=6, pady=6)
        ttk.Button(self.lf_out, text=tr("Browse…"), command=self.choose_out_dir).pack(side="left", padx=6, pady=6)

        # Transcription params
        self.lf_params = ttk.LabelFrame(self, text=tr("Transcription")); self.lf_params.pack(fill="x", padx=10, pady=8)

        self.lbl_model = ttk.Label(self.lf_params, text=tr("Model"))
        self.lbl_model.grid(row=0, column=0, sticky="w", padx=6, pady=6)
        self.model_var = tk.StringVar(value="small")
        ttk.Combobox(self.lf_params, textvariable=self.model_var, values=["tiny","base","small","medium","large"], state="readonly").grid(row=0, column=1, sticky="w", padx=6, pady=6)
        help_btn(self.lf_params, "help_model").grid(row=0, column=2, padx=4)

        self.lbl_chunk = ttk.Label(self.lf_params, text=tr("Chunk length (s)"))
        self.lbl_chunk.grid(row=0, column=3, sticky="w", padx=6, pady=6)
        self.chunk_var = tk.IntVar(value=30)
        ttk.Spinbox(self.lf_params, from_=5, to=600, textvariable=self.chunk_var, width=8).grid(row=0, column=4, sticky="w", padx=6, pady=6)
        help_btn(self.lf_params, "help_chunk").grid(row=0, column=5, padx=4)

        self.lbl_saveevery = ttk.Label(self.lf_params, text=tr("Save every N chunks"))
        self.lbl_saveevery.grid(row=1, column=0, sticky="w", padx=6, pady=6)
        self.save_every_var = tk.IntVar(value=5)
        ttk.Spinbox(self.lf_params, from_=1, to=100, textvariable=self.save_every_var, width=8).grid(row=1, column=1, sticky="w", padx=6, pady=6)
        help_btn(self.lf_params, "help_save_every").grid(row=1, column=2, padx=4)

        self.lbl_language = ttk.Label(self.lf_params, text=tr("Language (blank=auto)"))
        self.lbl_language.grid(row=1, column=3, sticky="w", padx=6, pady=6)
        self.lang_rec_var = tk.StringVar(value="en")
        ttk.Entry(self.lf_params, textvariable=self.lang_rec_var, width=10).grid(row=1, column=4, sticky="w", padx=6, pady=6)
        help_btn(self.lf_params, "help_language").grid(row=1, column=5, padx=4)

        self.lbl_known = ttk.Label(self.lf_params, text=tr("Known speakers (0=auto, 1=single)"))
        self.lbl_known.grid(row=2, column=0, sticky="w", padx=6, pady=6)
        self.known_spk_var = tk.IntVar(value=1)
        ttk.Spinbox(self.lf_params, from_=0, to=8, textvariable=self.known_spk_var, width=8).grid(row=2, column=1, sticky="w", padx=6, pady=6)
        help_btn(self.lf_params, "help_known_speakers").grid(row=2, column=2, padx=4)

        self.lbl_silence = ttk.Label(self.lf_params, text=tr("Silence → new paragraph (s)"))
        self.lbl_silence.grid(row=2, column=3, sticky="w", padx=6, pady=6)
        self.silence_var = tk.DoubleVar(value=2.0)
        ttk.Spinbox(self.lf_params, from_=0.5, to=10.0, increment=0.5, textvariable=self.silence_var, width=8).grid(row=2, column=4, sticky="w", padx=6, pady=6)
        help_btn(self.lf_params, "help_silence_gap").grid(row=2, column=5, padx=4)

        self.pre_clean_var = tk.BooleanVar(value=False)
        self.chk_preclean = ttk.Checkbutton(self.lf_params, text=tr("Pre-clean audio (denoise)"), variable=self.pre_clean_var)
        self.chk_preclean.grid(row=3, column=0, columnspan=2, sticky="w", padx=6, pady=6)
        help_btn(self.lf_params, "help_preclean").grid(row=3, column=2, padx=4)

        # Segments
        self.lf_segments = ttk.LabelFrame(self, text=tr("Segments")); self.lf_segments.pack(fill="x", padx=10, pady=8)
        self.lbl_maxseg = ttk.Label(self.lf_segments, text=tr("Max segment length (s, 0=off)"))
        self.lbl_maxseg.grid(row=0, column=0, sticky="w", padx=6, pady=6)
        self.max_seg_var = tk.IntVar(value=25)
        ttk.Spinbox(self.lf_segments, from_=0, to=180, textvariable=self.max_seg_var, width=8).grid(row=0, column=1, sticky="w", padx=6, pady=6)
        help_btn(self.lf_segments, "help_max_segment").grid(row=0, column=2, padx=4)

        # Highlight & markers
        self.lf_highlight = ttk.LabelFrame(self, text=tr("Concept Highlighter & Topic Markers")); self.lf_highlight.pack(fill="x", padx=10, pady=8)
        self.enable_kw_var = tk.BooleanVar(value=False)
        self.chk_kw = ttk.Checkbutton(self.lf_highlight, text=tr("Enable keyword highlighting + glossary"), variable=self.enable_kw_var)
        self.chk_kw.grid(row=0, column=0, sticky="w", padx=6, pady=6)
        help_btn(self.lf_highlight, "help_highlight").grid(row=0, column=1, padx=4)

        self.lbl_topn = ttk.Label(self.lf_highlight, text=tr("Top-N keywords"))
        self.lbl_topn.grid(row=0, column=2, sticky="w", padx=6, pady=6)
        self.topn_kw_var = tk.IntVar(value=25)
        ttk.Spinbox(self.lf_highlight, from_=5, to=100, textvariable=self.topn_kw_var, width=8).grid(row=0, column=3, sticky="w", padx=6, pady=6)

        self.lbl_markers = ttk.Label(self.lf_highlight, text=tr("Topic markers (comma-separated)"))
        self.lbl_markers.grid(row=1, column=0, sticky="w", padx=6, pady=6)
        self.markers_var = tk.StringVar(value="slide, theorem, lemma, proof, example, definition, chapter, section")
        ttk.Entry(self.lf_highlight, textvariable=self.markers_var, width=60).grid(row=1, column=1, columnspan=2, sticky="we", padx=6, pady=6)
        help_btn(self.lf_highlight, "help_markers").grid(row=1, column=3, padx=4)

        # Exports
        self.lf_outputs = ttk.LabelFrame(self, text=tr("Optional exports (raw .txt is always saved)")); self.lf_outputs.pack(fill="x", padx=10, pady=8)
        self.out_txt_refined = tk.BooleanVar(value=False)
        self.out_srt = tk.BooleanVar(value=False)
        self.out_vtt = tk.BooleanVar(value=False)
        self.out_md = tk.BooleanVar(value=False)
        self.out_docx = tk.BooleanVar(value=False)
        self.out_csvqa = tk.BooleanVar(value=False)
        self.out_llm_json = tk.BooleanVar(value=False)

        def add_option(row, col, label, var, help_key):
            chk = ttk.Checkbutton(self.lf_outputs, text=tr(label), variable=var)
            chk.grid(row=row, column=col, sticky="w", padx=6, pady=4)
            help_btn(self.lf_outputs, f"help_{help_key}").grid(row=row, column=col+1, padx=2)

        add_option(0, 0, "TXT refined", self.out_txt_refined, "txt_refined")
        add_option(0, 2, "SRT", self.out_srt, "srt")
        add_option(0, 4, "VTT", self.out_vtt, "vtt")
        add_option(1, 0, "Markdown", self.out_md, "md")
        add_option(1, 2, "DOCX", self.out_docx, "docx")
        add_option(1, 4, "CSV Q&A", self.out_csvqa, "csvqa")
        add_option(1, 6, "JSON for LLM", self.out_llm_json, "llm_json")

        # Run section
        fr_run = ttk.Frame(self); fr_run.pack(fill="x", padx=10, pady=8)
        self.btn_start = ttk.Button(fr_run, text=tr("Start"), command=self.start_work); self.btn_start.pack(side="left", padx=6)
        self.progress = ttk.Progressbar(fr_run, orient="horizontal", mode="determinate"); self.progress.pack(side="left", fill="x", expand=True, padx=10)
        self.lbl_progress = ttk.Label(fr_run, text=tr("Idle")); self.lbl_progress.pack(side="left", padx=6)

        # Log
        self.lf_log = ttk.LabelFrame(self, text="Log"); self.lf_log.pack(fill="both", expand=True, padx=10, pady=8)
        self.txt_log = tk.Text(self.lf_log, height=12, wrap="word"); self.txt_log.pack(fill="both", expand=True, padx=6, pady=6)

        # Start draining log
        self.after(100, self.drain_log_queue)

    # ---------- language switching ----------
    def on_change_lang(self):
        global _current_lang
        new_lang = self.lang_var.get()
        if new_lang not in LANGS: return
        _current_lang = new_lang
        self.ui_lang = new_lang
        save_ui_lang(new_lang)
        self.refresh_texts()

    def refresh_texts(self):
        self.title(tr("APP_NAME"))
        self.lf_files.config(text=tr("Media files"))
        self.btn_add.config(text=tr("Add files…"))
        self.btn_clear.config(text=tr("Clear"))
        if not getattr(self, "files", []):
            self.lbl_files.config(text=tr("No files selected"))
        self.lf_out.config(text=tr("Output folder"))
        self.lf_params.config(text=tr("Transcription"))
        self.lbl_model.config(text=tr("Model"))
        self.lbl_chunk.config(text=tr("Chunk length (s)"))
        self.lbl_saveevery.config(text=tr("Save every N chunks"))
        self.lbl_language.config(text=tr("Language (blank=auto)"))
        self.lbl_known.config(text=tr("Known speakers (0=auto, 1=single)"))
        self.lbl_silence.config(text=tr("Silence → new paragraph (s)"))
        self.chk_preclean.config(text=tr("Pre-clean audio (denoise)"))
        self.lf_segments.config(text=tr("Segments"))
        self.lbl_maxseg.config(text=tr("Max segment length (s, 0=off)"))
        self.lf_highlight.config(text=tr("Concept Highlighter & Topic Markers"))
        self.chk_kw.config(text=tr("Enable keyword highlighting + glossary"))
        self.lbl_topn.config(text=tr("Top-N keywords"))
        self.lbl_markers.config(text=tr("Topic markers (comma-separated)"))
        self.lf_outputs.config(text=tr("Optional exports (raw .txt is always saved)"))
        self.btn_start.config(text=tr("Start"))
        self.lbl_progress.config(text=tr("Idle"))

    # ---------- file ops ----------
    def add_files(self):
        paths = filedialog.askopenfilenames(title="Select media files", filetypes=[("Video/Audio","*.mp4 *.mkv *.mov *.mp3 *.wav *.m4a"), ("All files","*.*")])
        if paths:
            self.files.extend(paths); self.update_files_label()

    def clear_files(self):
        self.files = []; self.update_files_label()

    def update_files_label(self):
        if not self.files: self.lbl_files.config(text=tr("No files selected")); return
        shown = "; ".join([os.path.basename(p) for p in self.files[:3]])
        if len(self.files) > 3: shown += f" … (+{len(self.files)-3})"
        self.lbl_files.config(text=shown)

    def choose_out_dir(self):
        d = filedialog.askdirectory(title="Choose output folder")
        if d: self.out_dir = d; self.out_var.set(d)

    # ---------- logging / progress ----------
    def log(self, msg): self.log_queue.put(msg)

    def drain_log_queue(self):
        try:
            while True:
                msg = self.log_queue.get_nowait()
                self.txt_log.insert("end", msg + "\n"); self.txt_log.see("end")
        except queue.Empty:
            pass
        self.after(100, self.drain_log_queue)

    def set_progress(self, fraction, elapsed, eta):
        self.progress["value"] = int(fraction*100)
        def fmt(t):
            if t >= 3600: h=int(t//3600); m=int((t%3600)//60); return f"{h}h {m}m"
            if t >= 60: m=int(t//60); s=int(t%60); return f"{m}m {s}s"
            return f"{int(t)}s"
        self.lbl_progress.config(text=f"{int(fraction*100)}% | elapsed {fmt(elapsed)} | ETA {fmt(eta)}")
        self.update_idletasks()

    # ---------- run ----------
    def start_work(self):
        if hasattr(self, "worker_thread") and self.worker_thread and self.worker_thread.is_alive():
            messagebox.showinfo(tr("APP_NAME"), tr("A job is already running."))
            return
        if not hasattr(self, "files") or not self.files:
            messagebox.showwarning(tr("APP_NAME"), tr("Select at least one media file."))
            return
        if not check_ffmpeg():
            messagebox.showerror(tr("APP_NAME"), tr("ffmpeg not found in PATH."))
            return

        model_name = self.model_var.get().strip()
        chunk_len = max(1, int(self.chunk_var.get()))
        save_every = max(1, int(self.save_every_var.get()))
        language = self.lang_rec_var.get().strip() or None
        known_speakers = int(self.known_spk_var.get())
        silence_gap = float(self.silence_var.get())
        out_dir = self.out_var.get().strip() or self.out_dir
        os.makedirs(out_dir, exist_ok=True)
        pre_clean_audio = bool(self.pre_clean_var.get())
        max_segment_seconds = int(self.max_seg_var.get())
        enable_keywords = bool(self.enable_kw_var.get())
        topn_keywords = int(self.topn_kw_var.get())
        markers_list = self.markers_var.get()
        export_txt_refined = bool(self.out_txt_refined.get())
        export_srt = bool(self.out_srt.get())
        export_vtt = bool(self.out_vtt.get())
        export_md = bool(self.out_md.get())
        export_docx = bool(self.out_docx.get())
        export_csvqa = bool(self.out_csvqa.get())
        export_llm_json = bool(self.out_llm_json.get())

        self.btn_start.config(state="disabled"); self.progress["value"]=0
        self.lbl_progress.config(text="Loading model…"); self.txt_log.delete("1.0","end")

        def work():
            try:
                self.log(f"Loading Whisper model '{model_name}'… (first run may download weights)")
                model = whisper.load_model(model_name)
                for i, media in enumerate(self.files, 1):
                    self.log(f"[{i}/{len(self.files)}] {media}")
                    transcribe_one(
                        model=model,
                        media_path=media,
                        chunk_length=chunk_len,
                        save_every_chunks=save_every,
                        language=language,
                        known_speakers=known_speakers,
                        silence_gap=silence_gap,
                        out_dir=out_dir,
                        pre_clean_audio=pre_clean_audio,
                        max_segment_seconds=max_segment_seconds,
                        enable_keywords=enable_keywords,
                        topn_keywords=topn_keywords,
                        markers_list=markers_list,
                        export_txt_refined=export_txt_refined,
                        export_srt=export_srt,
                        export_vtt=export_vtt,
                        export_md=export_md,
                        export_docx=export_docx,
                        export_csvqa=export_csvqa,
                        export_llm_json=export_llm_json,
                        progress_cb=self.set_progress,
                        log=self.log
                    )
                notification.notify(title=tr("APP_NAME"), message="All transcriptions completed!", app_name=tr("APP_NAME"), timeout=8)
                self.log("Done.")
            except Exception as e:
                self.log(f"ERROR: {e}"); messagebox.showerror(tr("APP_NAME"), str(e))
            finally:
                self.btn_start.config(state="normal"); self.lbl_progress.config(text=tr("Idle")); self.progress["value"]=0

        self.worker_thread = threading.Thread(target=work, daemon=True); self.worker_thread.start()

if __name__ == "__main__":
    App().mainloop()
